from django.db import models
import docx


class Article(models.Model):
    title = models.CharField(max_length=100)
    content = models.TextField(blank=True, null=True)
    image = models.ImageField(upload_to='articles/', blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)
    document = models.FileField(upload_to='articles/docs/', blank=True, null=True)

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)

        if self.document and (not self.content or not self.content.strip()):
            try:
                doc = docx.Document(self.document.path)
                parts = []

            
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)

               
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            parts.append(row_text)

                self.content = "\n".join(parts)
                
                super().save(update_fields=["content"])
            except Exception as e:
                print("Error extracting docx:", e)

    def __str__(self):
        return self.title

class Event(models.Model):
    title = models.CharField(max_length=200)
    description = models.TextField()
    tag = models.CharField(max_length=50, blank=True, null=True)
    image = models.ImageField(upload_to="about/", blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = "Event"            
        verbose_name_plural = "Events"

    def __str__(self):
        return self.title


    
